from docx import Document
from lxml import etree
from zipfile import ZipFile
import mysql.connector
from fastapi import FastAPI, UploadFile, HTTPException
import os
import json
import re

# Configuração do Banco de Dados
def get_db_connection():
    return mysql.connector.connect(
        host="localhost",
        user="root",
        password="",
        # database="livro_insert"
        database="livro_crud"
    )

def execute_insert(sql, params):
    try:
        db = get_db_connection()
        cursor = db.cursor()
        cursor.execute(sql, params)
        last_id = cursor.lastrowid
        db.commit()
        return last_id
    except Exception as e:
        print(f"Erro ao executar o SQL: {sql}, Parâmetros: {params}, Erro: {e}")
        raise
    finally:
        cursor.close()
        db.close()

def get_titulo_id(livro_id, conteudo):
    sql = "SELECT id FROM titulos WHERE livro_id = %s AND conteudo = %s"
    try:
        db = get_db_connection()
        cursor = db.cursor()
        cursor.execute(sql, (livro_id, conteudo))
        result = cursor.fetchone()
        return result[0] if result else None
    finally:
        cursor.close()
        db.close()

def get_capitulo_id(titulo_id, conteudo):
    sql = "SELECT id FROM capitulos WHERE titulo_id = %s AND conteudo = %s"
    try:
        db = get_db_connection()
        cursor = db.cursor()
        cursor.execute(sql, (titulo_id, conteudo))
        result = cursor.fetchone()
        return result[0] if result else None
    finally:
        cursor.close()
        db.close()

def get_secao_id(capitulo_id, conteudo):
    sql = "SELECT id FROM secaos WHERE capitulo_id = %s AND conteudo = %s"
    try:
        db = get_db_connection()
        cursor = db.cursor()
        cursor.execute(sql, (capitulo_id, conteudo))
        result = cursor.fetchone()
        return result[0] if result else None
    finally:
        cursor.close()
        db.close()

# Funções de Inserção no Banco
def insert_livro(conteudo):
    sql = "INSERT INTO livros (conteudo) VALUES (%s)"
    return execute_insert(sql, (conteudo,))

def insert_titulo(livro_id, conteudo):
    titulo_id = get_titulo_id(livro_id, conteudo)
    if not titulo_id:
        titulo_id = execute_insert("INSERT INTO titulos (livro_id, conteudo) VALUES (%s, %s)", (livro_id, conteudo))
        print(f"Título criado: {conteudo} (ID: {titulo_id})")
    else:
        print(f"Título já existente: {conteudo} (ID: {titulo_id})")
    return titulo_id

def insert_capitulo(titulo_id, conteudo):
    capitulo_id = get_capitulo_id(titulo_id, conteudo)
    if not capitulo_id:
        capitulo_id = execute_insert("INSERT INTO capitulos (titulo_id, conteudo) VALUES (%s, %s)", (titulo_id, conteudo))
        print(f"Capítulo criado: {conteudo} (ID: {capitulo_id})")
    else:
        print(f"Capítulo já existente: {conteudo} (ID: {capitulo_id})")
    return capitulo_id

def insert_secao(capitulo_id, conteudo):
    secao_id = get_secao_id(capitulo_id, conteudo)
    if not secao_id:
        secao_id = execute_insert("INSERT INTO secaos (capitulo_id, conteudo) VALUES (%s, %s)", (capitulo_id, conteudo))
        print(f"Seção criada: {conteudo} (ID: {secao_id})")
    else:
        print(f"Seção já existente: {conteudo} (ID: {secao_id})")
    return secao_id

def insert_artigo(secao_id, conteudo):
    sql = "INSERT INTO artigos (secao_id, conteudo) VALUES (%s, %s)"
    return execute_insert(sql, (secao_id, conteudo))

def insert_paragrafo(artigo_id, conteudo, tipo=None):
    sql = "INSERT INTO paragrafos (artigo_id, conteudo, tipo) VALUES (%s, %s, %s)"
    return execute_insert(sql, (artigo_id, conteudo, tipo))

def insert_quadro(associado_id, tipo, conteudo):
    sql = "INSERT INTO quadros_esquemas (associado_id, tipo, conteudo) VALUES (%s, %s, %s)"
    return execute_insert(sql, (associado_id, tipo, json.dumps(conteudo)))

def insert_nota_rodape(associado_a, associado_id, conteudo):
    sql = "INSERT INTO nota_rodapes (associado_a, associado_id, conteudo) VALUES (%s, %s, %s)"
    return execute_insert(sql, (associado_a, associado_id, conteudo))

def insert_remissao(paragrafo_id, conteudo):
    sql = "INSERT INTO remissaos (paragrafo_de_id, conteudo) VALUES (%s, %s)"
    return execute_insert(sql, (paragrafo_id, conteudo))

def process_titulos(livro_id, conteudo):
    return insert_titulo(livro_id, conteudo)

def process_capitulos(titulo_id, conteudo):
    return insert_capitulo(titulo_id, conteudo)

def process_artigos(secao_id, conteudo):
    return insert_artigo(secao_id, conteudo)

def process_paragrafos(artigo_id, conteudo, tipo=None):
    return insert_paragrafo(artigo_id, conteudo, tipo)

# Extração de Notas de Rodapé
def extract_notes(docx_path, note_type="footnotes"):
    notes = []
    with ZipFile(docx_path, 'r') as docx:
        xml_path = f"word/{note_type}.xml"
        if xml_path in docx.namelist():
            with docx.open(xml_path) as xml_file:
                xml_content = xml_file.read()
                tree = etree.fromstring(xml_content)
                namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                for note in tree.xpath("//w:footnote | //w:endnote", namespaces=namespace):
                    note_id = note.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id")
                    content = "".join(note.xpath(".//w:t/text()", namespaces=namespace))
                    notes.append({"id": note_id, "content": content})
    return notes

def extract_tables_from_docx(doc):
    tables_data = []
    for table in doc.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        rows = [
            [cell.text.strip() for cell in row.cells]
            for row in table.rows[1:]
        ]
        tables_data.append({
            "header": headers,
            "rows": rows
        })
    return tables_data

# Processar o Documento com Marcações ###
def process_document(file_path):
    document = Document(file_path)
    estrutura = {"elementos": [], "tabelas": [], "notas_rodape": {}}

    # Extraindo notas de rodapé
    notas = extract_notes(file_path, "footnotes")
    for nota in notas:
        estrutura["notas_rodape"][f"###nota {nota['id']}###"] = nota["content"]

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        match = re.match(r"###(.+?)###", text)
        if match:
            tipo = match.group(1).strip()
            estrutura["elementos"].append({"tipo": tipo, "conteudo": ""})
        else:
            if estrutura["elementos"]:
                estrutura["elementos"][-1]["conteudo"] += text + "\n"

    estrutura["tabelas"] = extract_tables_from_docx(document)

    return estrutura

def processar_livro(file_path):
    estrutura = process_document(file_path)
    livro_id = insert_livro("Regimento Interno Comentado")
    print(f"Livro inserido com ID: {livro_id}")
    titulo_id = insert_titulo(livro_id, "Título Geral")
    capitulo_id = insert_capitulo(titulo_id, "Capítulo Geral")
    secao_id = insert_secao(capitulo_id, "Seção Geral")
    artigo_id = insert_artigo(secao_id, "Artigo Geral")
    print(f"Hierarquia inicial criada: Título {titulo_id}, Capítulo {capitulo_id}, Seção {secao_id}, Artigo {artigo_id}")

    for elemento in estrutura["elementos"]:
        tipo = elemento["tipo"]
        conteudo = elemento["conteudo"].strip()

        if not conteudo:
            continue

        for nota_ref, nota_conteudo in estrutura["notas_rodape"].items():
            if nota_ref in conteudo:
                conteudo = conteudo.replace(nota_ref, "") 
                insert_nota_rodape(tipo, livro_id, nota_conteudo)
                print(f"Nota de rodapé associada: {nota_ref} -> {nota_conteudo}")

        if tipo == "titulos":
            titulo_id = process_titulos(livro_id, conteudo)
            print(f"Título inserido: {conteudo}")
            capitulo_id = insert_capitulo(titulo_id, "Capítulo Geral")
            secao_id = insert_secao(capitulo_id, "Seção Geral")
            artigo_id = insert_artigo(secao_id, "Artigo Geral")
        elif tipo == "capitulos":
            capitulo_id = process_capitulos(titulo_id, conteudo)
            print(f"Capítulo inserido: {conteudo}")
            secao_id = insert_secao(capitulo_id, "Seção Geral")
            artigo_id = insert_artigo(secao_id, "Artigo Geral")
        elif tipo == "secaos":
            secao_id = insert_secao(capitulo_id, conteudo)
            print(f"Seção inserida: {conteudo}")
            artigo_id = insert_artigo(secao_id, "Artigo Geral")
        elif tipo == "artigos":
            artigo_id = process_artigos(secao_id, conteudo)
            print(f"Artigo inserido: {conteudo}")
        elif tipo.startswith("paragrafo"):
            tipo_paragrafo = re.search(r"tipo: (.+)", tipo)
            tipo_paragrafo = tipo_paragrafo.group(1) if tipo_paragrafo else None
            process_paragrafos(artigo_id, conteudo, tipo_paragrafo)
            print(f"Parágrafo inserido: {conteudo}")
        elif tipo == "remissaos":
            if artigo_id is None:
                raise ValueError("Erro: Nenhum artigo associado para a remissão.")
            insert_remissao(artigo_id, conteudo)
            print(f"Remissão inserida: {conteudo}")
        else:
            if artigo_id is None:
                if secao_id is None:
                    secao_id = insert_secao(capitulo_id, "Seção Geral")
                    print("Seção Geral criada.")
                if capitulo_id is None:
                    capitulo_id = insert_capitulo(titulo_id, "Capítulo Geral")
                    print("Capítulo Geral criado.")
                if titulo_id is None:
                    titulo_id = insert_titulo(livro_id, "Título Geral")
                    print("Título Geral criado.")
                artigo_id = insert_artigo(secao_id, "Artigo Geral")
                print("Artigo Geral criado.")

    for tabela in estrutura["tabelas"]:
        insert_quadro(livro_id, "tabela", tabela)
        print(f"Tabela inserida: {tabela['header']}")

# API com FastAPI
app = FastAPI()

@app.post("/upload")
async def upload_file(file: UploadFile):
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Formato de arquivo inválido. Somente arquivos .docx são suportados.")
    temp_file_path = f"./temp_{file.filename}"
    try:
        with open(temp_file_path, "wb") as temp_file:
            temp_file.write(await file.read())
        processar_livro(temp_file_path)
        return {"status": "success", "message": f"Arquivo '{file.filename}' processado com sucesso."}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro ao processar o arquivo: {e}")
    finally:
        if os.path.exists(temp_file_path):
            os.remove(temp_file_path)

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8002)